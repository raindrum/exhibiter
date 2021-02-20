# Exhibiter, copyright (c) 2021 Simon Raindrum Sherred.
# This software may not be used to evict people, see LICENSE.md.

# python standard imports
from re import search, sub, fullmatch
from pathlib import Path

# third-party imports
from pdfrw import PdfReader, PdfWriter, buildxobj, toreportlab
from reportlab.lib import pagesizes, colors
from reportlab.pdfgen.canvas import Canvas
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx import Document
from PIL import Image

# global variables
FILE_TYPES = ["png", "PNG", "jpg", "JPG", "jpeg", "JPEG", "pdf", "PDF"]
EXCLUDE_PATTERN = "\((UNUSED|[Uu]nused)\)"
DISPUTE_FILE = "evidentiary disputes.txt"


class Exhibit:
    """
    The the heart of Exhibiter. This object contains a ReportLab
    canvas and a list of documents, corresponding to the contents of
    one exhibit.
    """

    @classmethod
    def from_path(
        cls,
        exhibit_path: Path,
        respect_exclusions: bool = True,
        number_pages: bool = True,
        page_label_coords: tuple = (50, 3),
        rotate_landscape_pics: bool = True,
        strip_leading_digits: bool = True,
    ):
        """
        This constructor makes an exhibit from a given folder or file.

        The folder name should be formatted like this: "101" or
        "101. Residential Lease". This will set the exhibit's index
        and, optionally, its title.

        The exhibit folder should contain one or more documents. A
        document means a file in the FILE_TYPES list, or a folder full
        of such files. Documents whose names contain "(UNUSED)" will
        normally be omitted."""

        # throw error if filename is wrong
        if not fullmatch("^(\d+|[A-Y])(\.?( .+)?)?", exhibit_path.stem):
            raise SyntaxError(
                f"'{exhibit_path.stem}' isnt a valid name for an exhibit. It"
                + " must be a number or capital letter from A-Y, optionally"
                + " followed by a title to display in the exhibit list. Valid"
                + " examples include names like these:"
                + '\n"101"'
                + '\n"102. Party Communications"'
                + '\n"A"'
            )

        # for single-document exhibits, throw errors if they're the
        # wrong type, or if they don't have titles.
        if not exhibit_path.is_dir():
            if exhibit_path.suffix[1:] not in FILE_TYPES:
                raise SyntaxError(
                    f"{exhibit_path.name} is not a supported file type."
                    + " Exhibits can be PDFs, JPGs, PNGs, or folders"
                    + " full of those things."
                )
            elif not fullmatch("^(\d+|[A-Y])\. .+", exhibit_path.stem):
                raise SyntaxError(
                    f'"{exhibit_path.name}" is not a valid name for a'
                    + ' single-document exhibit. It must have a title,'
                    + ' like "101. Rental Agreement.pdf".'
                )

        # get index and title (if any) from filename.
        folder_name = exhibit_path.name
        folder_name = _process_filename(exhibit_path.name, False)
        sections = folder_name.split(". ", 1)
        index = sections[0]
        
        # add a title only if the exhibit path is a directory. For one-file
        # exhibits, the document name makes a title unnecessary
        if len(sections) > 1 and exhibit_path.is_dir():
            title = sections[1]
        else:
            title = None
        
        # get evidentiary disputes
        dispute_file = exhibit_path / DISPUTE_FILE
        if dispute_file.exists():
            evidentiary_disputes = dispute_file.read_text()
        else:
            evidentiary_disputes = None
        # generate exhibit
        exhibit = cls(
            index,
            title,
            number_pages=number_pages,
            page_label_coords=page_label_coords,
            rotate_landscape_pics=rotate_landscape_pics,
            evidentiary_disputes=evidentiary_disputes,
        )

        # add all evidence from the path to it
        if exhibit_path.is_dir():
            for path in evidence_in_dir(exhibit_path, respect_exclusions):
                exhibit.add_doc(path, strip_leading_digits=strip_leading_digits)
        else:  # if the exhibit is just a single file, add it
            exhibit.add_doc(
                exhibit_path,
                title=sub("^(\d+|[A-Z])\. ", "", exhibit_path.stem)
            )
        
        return exhibit

    def __init__(
        self,
        index: str,
        title: str = None,
        number_pages: bool = True,
        page_label_coords: tuple = (50, 3),
        rotate_landscape_pics: bool = True,
        evidentiary_disputes: str = None,
    ):
        """
        This creates a bare-bones exhibit with only a cover sheet.
        You can then populate it by running add_doc() one or more times.
        """

        # make a canvas write a cover page like "EXHIBIT 101"
        canvas = Canvas("never_save_to_this_path.pdf")
        canvas.setPageSize(pagesizes.letter)
        canvas.setFont("Helvetica", 32)
        x, y = canvas._pagesize[0] / 2, canvas._pagesize[1] / 7
        canvas.drawCentredString(x, y, f"EXHIBIT {index}")
        canvas.showPage()

        # set this exhibit's various variables
        self.canvas: Canvas = canvas
        self.documents: list = []
        self.index: str = index
        self.title: str = title
        self.evidentiary_disputes: str = evidentiary_disputes
        self.number_pages: bool = number_pages
        self.rotate_landscape_pics: bool = rotate_landscape_pics
        self.page_label_coords: tuple = page_label_coords
        self.page_count: int = 0

    def add_doc(
        self,
        doc_path: Path,
        respect_exclusions: bool = True,
        strip_leading_digits: bool = True,
        title: str = None,
    ):
        """
        Adds a document (i.e. an image, PDF, or a folder of either)
        to this exhibit.
        """
        startpage = self.page_count + 1

        if not title:
            title = _process_filename(doc_path.stem, strip_leading_digits)

        if doc_path.is_dir():  # walk through directory and add files to doc
            file_paths = []
            for extension in FILE_TYPES:
                file_paths += doc_path.glob("**/*." + extension)

            # add each file in order, except the ones marked for omission
            for path in sorted(file_paths):
                if respect_exclusions and search(EXCLUDE_PATTERN, path.stem):
                    continue
                self._insert_pdf_or_image(path)

        else: # add single-file document to exhibit
            self._insert_pdf_or_image(doc_path)
        self.documents.append(
            {"name": title, "page_span": (startpage, self.page_count), "path": doc_path}
        )

    def _insert_pdf_or_image(self, path: Path):
        """
        Checks whether the given file is a PDF or an image, and
        performs the appropriate actions to add it to the main PDF.
        """

        if path.suffix in [".pdf", ".PDF"]:
            pages = PdfReader(path).pages
            pages = [buildxobj.pagexobj(page) for page in pages]
            for page in pages:
                self.canvas.setPageSize((page.BBox[2], page.BBox[3]))
                self.canvas.doForm(toreportlab.makerl(self.canvas, page))
                self._finish_page()

        elif path.suffix[1:] in FILE_TYPES:  # treat path as an image
            self.canvas.setPageSize(pagesizes.letter)
            page_w, page_h = self.canvas._pagesize
            # Rotate landscape images to fit portrait page
            img = Image.open(path)
            img_ratio = img.size[0] / img.size[1]
            if img_ratio > 1 and self.rotate_landscape_pics:
                self.canvas.saveState()
                self.canvas.rotate(-90)
                w, h = 0.9 * page_h, 0.9 * page_w
                rotated = True
                x = (-w - page_h) / 2
                y = (page_w - h) / 2
                # x = (page_h - w) / 2
                # y = (-page_w - h) / 2
            else:
                w, h = 0.9 * page_w, 0.9 * page_h
                x = (page_w - w) / 2
                y = (page_h - h) / 2
                rotated = False
            self.canvas.drawImage(path, x, y, w, h, preserveAspectRatio=True)
            if rotated:
                self.canvas.restoreState()
            self._finish_page()

        else:
            raise SyntaxError(f"{path} is not a supported type: {FILE_TYPES}")

    def _finish_page(self):
        """Print a page number (maybe), then move on to the next page."""
        self.page_count += 1
        if self.number_pages:
            string = f"{self.index}-{self.page_count}"
            mid = [
                self.canvas._pagesize[x] * self.page_label_coords[x] / 100
                for x in [0, 1]
            ]
            self.canvas.setFillColor(colors.white)
            self.canvas.rect(mid[0] - 25, mid[1] - 4, 50, 15, stroke=0, fill=1)
            self.canvas.setFillColor(colors.black)
            self.canvas.drawCentredString(mid[0], mid[1], string)
        self.canvas.showPage()

    def __str__(self):
        return self.path.stem


# ######################################################################
# Module-Level Functions
# ######################################################################


def write_pdf(exhibits: list[Exhibit], output_path: str):
    """Save the given list of exhibits to a PDF document."""
    writer = PdfWriter()
    for exhibit in exhibits:
        reader = PdfReader(fdata=exhibit.canvas.getpdfdata())
        writer.addpages(reader.pages)
    writer.write(output_path)


def write_list(
    exhibits: list[Exhibit],
    output_path: str,
    attachment_no: int = 4,
    party_label: str = "Defense",
    show_page_numbers: bool = True,
    reserve_rebuttal: bool = True,
):
    """Save a Word document listing the given exhibits in detail."""
    template = str(Path(__file__).parent.absolute() / "template.docx")
    exhibit_list = Document(template)

    # fill in title line
    exhibit_list.paragraphs[0].add_run(
        f"Attachment {attachment_no}: Exhibit List"
    ).bold = True
    
    # write table header
    exhibit_list.tables[0].rows[0].cells[0].paragraphs[0].add_run(
        f"{party_label.upper()} EXHIBITS"
    ).bold = True
    
    for exhibit in exhibits:
        # add a table row, to represent the exhibit
        row = exhibit_list.tables[0].add_row()
        row.cells[0].text = exhibit.index
        
        # center-align the first two cells
        for c in [0, 1]:  # center-align columns 0 and 1
            row.cells[c].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # write the exhibit title
        if exhibit.title:
            row.cells[3].paragraphs[0].text = exhibit.title + ":"
        
        # add a line for each document in the exhibit
        for i, doc in enumerate(exhibit.documents):
            # use an existing blank paragraph then make one for each doc
            if i == 0 and not exhibit.title:
                paragraph = row.cells[3].paragraphs[0]
            else:
                paragraph = row.cells[3].add_paragraph()
            description = doc["name"]
            if i > 0 and show_page_numbers:  # e.g. "(p.32)"
                description += f" (p.{doc['page_span'][0]})"
            paragraph.text = description
        row.cells[4].text = exhibit.evidentiary_disputes or ""

    if reserve_rebuttal:
        # calculate the next exhibit number or letter
        last_index = exhibits[-1].index
        if search("[A-Y]", last_index):
            next_index = chr(ord(last_index) + 1)
        else:
            next_index = str(int(last_index) + 1)
        
        # reserve that exhibit for rebuttal
        row = exhibit_list.tables[0].add_row()
        row.cells[0].text = next_index
        row.cells[3].text = "Reserved for Rebuttal"
        for c in [0, 1]:
            row.cells[c].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

    exhibit_list.save(output_path)


def evidence_in_dir(folder: Path, respect_exclusions: bool = True):
    returns = []
    for path in sorted(folder.iterdir()):
        # skip unsupported files
        if not path.is_dir() and path.suffix[1:] not in FILE_TYPES:
            continue
        # skip files containing the exclude pattern (by default)
        if respect_exclusions and search(EXCLUDE_PATTERN, path.stem):
            continue
        returns.append(path)
    if returns:
        return returns
    else:
        raise FileNotFoundError(f"{folder} doesn't seem to contain any evidence.")


def _process_filename(name: str, strip_leading_digits: bool = True) -> str:
    """
    Convert a filename into a document description.
    Strips EXCLUDE_PATTERN, if present, and moves leading YYYY-MM-DD
    dates to the end, in M/D/YY format. Also, converts things like
    "1. First Document" to "First Document" by default.
    """
    # remove file extension
    filetypes_regex = '(\.' + '|\.'.join(FILE_TYPES) + ')$'
    name = sub(filetypes_regex, '', name)
    
    # remove the exclude pattern
    name = sub(" ?(" + EXCLUDE_PATTERN + ")", "", name)
    
    # if name begins with YYYY-MM-DD, put M/D/YYYY at end
    if search("^\d{4}-\d{2}-\d{2} .+", name):
        name = (
            name[11:]
            + " "
            + name[5:7].lstrip("0")
            + "/"
            + name[8:10].lstrip("0")
            + "/"
            + name[2:4]
        )
    # if name starts with a number, period, and space, remove it
    elif strip_leading_digits:
        name = sub("^\d+\. ", "", name)
    return name
